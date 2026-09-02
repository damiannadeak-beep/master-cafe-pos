import pandas as pd
from docx import Document
from docx.shared import Pt
import os

questions = [
    {
        "No": 1,
        "Topik": "Kendala Operasional Harian",
        "Pertanyaan": "Apa kendala terbesar yang paling sering mengganggu kelancaran pesanan dari meja pelanggan hingga ke dapur?"
    },
    {
        "No": 2,
        "Topik": "Keuangan & Kasir (Shift)",
        "Pertanyaan": "Apakah pernah atau sering terjadi selisih antara total catatan penjualan dengan uang fisik yang ada di laci kasir saat tutup warung?"
    },
    {
        "No": 3,
        "Topik": "Pemantauan Stok Bahan Baku",
        "Pertanyaan": "Bagaimana cara Bapak/Ibu mengecek sisa bahan baku? Apakah karyawan pernah lupa mencatat sehingga bahan tiba-tiba habis saat jam sibuk?"
    },
    {
        "No": 4,
        "Topik": "Efisiensi Pembayaran",
        "Pertanyaan": "Bagaimana cara kasir mencatat dan memverifikasi pembayaran digital (seperti QRIS atau transfer)? Apakah prosesnya sudah cepat atau masih memakan waktu?"
    },
    {
        "No": 5,
        "Topik": "Solusi Pemesanan Mandiri (Self-Order)",
        "Pertanyaan": "Jika ada sistem di mana pelanggan bisa memesan dan langsung membayar sendiri dari meja (via Scan QR), apakah menurut Bapak/Ibu itu akan sangat membantu?"
    },
    {
        "No": 6,
        "Topik": "Harapan Terhadap Sistem Baru",
        "Pertanyaan": "Fitur atau jenis laporan seperti apa yang paling Bapak/Ibu butuhkan agar pengelolaan Master Cafe menjadi jauh lebih mudah ke depannya?"
    }
]

closing_statement = "Sebagai informasi, saat ini kami sedang mengembangkan Sistem Point of Sales (POS) modern berbasis web. Aplikasi ini telah dilengkapi dengan fitur pemesanan mandiri via QR Code (Self-Order), manajemen shift kasir untuk mencegah selisih uang, pemotongan stok bahan baku secara otomatis, serta integrasi pembayaran digital. Tujuan dari wawancara ini adalah untuk memahami kebutuhan spesifik Master Cafe Bengkalis, sehingga kami dapat menyesuaikan dan menyempurnakan fitur-fitur aplikasi ini agar benar-benar tepat sasaran dan mampu menjadi solusi nyata bagi operasional harian Master Cafe."

# 1. Create Word Document
doc = Document()
title = doc.add_heading('DAFTAR PERTANYAAN WAWANCARA', level=1)
title.alignment = 1 # Center

subtitle = doc.add_paragraph('Studi Kasus: Pemilik Master Cafe Bengkalis')
subtitle.alignment = 1 # Center

doc.add_paragraph('\n')

for q in questions:
    p = doc.add_paragraph()
    p.add_run(f"{q['No']}. {q['Topik']}\n").bold = True
    p.add_run(f"{q['Pertanyaan']}\n")

doc.add_paragraph('\n')
p_closing = doc.add_paragraph()
p_closing.add_run("Penutup & Latar Belakang Aplikasi:\n").bold = True
p_closing.add_run(closing_statement)

# Use a standard, formal font (Times New Roman, 12pt)
for paragraph in doc.paragraphs:
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(12)

docx_path = r'c:\xampp\htdocs\angkringan-pos\docs\Pertanyaan_Wawancara_Master_Cafe.docx'
doc.save(docx_path)
print(f"Created Word Document at {docx_path}")

# 2. Update Excel File
excel_path = r'c:\xampp\htdocs\angkringan-pos\docs\Pertanyaan_Wawancara_Master_Cafe.xlsx'
# Add closing statement as a final row
excel_data = questions.copy()
excel_data.append({
    "No": "",
    "Topik": "Penutup & Latar Belakang",
    "Pertanyaan": closing_statement
})

df = pd.DataFrame(excel_data)
df.to_excel(excel_path, index=False)
print(f"Updated Excel Document at {excel_path}")
